from flask import Flask, render_template, request, send_file, jsonify
from flask_socketio import SocketIO, emit
import io
import numpy as np
import threading
import os
import tempfile
import wave
from docx import Document
from docx.shared import Pt
from werkzeug.utils import secure_filename
import whisper
import torch

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your-secret-key-123'
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file size

# Ensure upload directory exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# File upload disabled - using Web Speech API only

socketio = SocketIO(app, cors_allowed_origins="*", async_mode='threading')

# Load Whisper model
print("Loading Whisper model...")
DEVICE = "cuda" if torch.cuda.is_available() else "cpu"
model = whisper.load_model("base", device=DEVICE)
USE_WHISPER = True
print("Whisper model loaded successfully")

# Store transcriptions
transcriptions = []

@app.route('/')
def index():
    return render_template('index.html')

@socketio.on('audio_chunk')
def handle_audio_chunk(data):
    """Handle incoming audio chunks"""
    global transcriptions

    try:
        # Convert audio chunk to numpy array
        audio_data = np.frombuffer(data, dtype=np.int16)

        # Save audio chunk to a temporary WAV file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as temp_audio:
            with wave.open(temp_audio, 'wb') as wf:
                wf.setnchannels(1)
                wf.setsampwidth(2)
                wf.setframerate(16000)
                wf.writeframes(audio_data.tobytes())

            # Transcribe audio file
            result = model.transcribe(temp_audio.name)

            # Append to transcriptions
            transcriptions.append(result['text'])

            # Send update to all clients
            emit('update', {
                'text': result['text'],
                'full_text': " ".join(transcriptions)
            }, broadcast=True)

            # Clean up temp file
            os.unlink(temp_audio.name)

    except Exception as e:
        print(f"Error processing audio chunk: {e}")

@socketio.on('check_whisper')
def check_whisper():
    """Check if Whisper is available"""
    emit('whisper_status', {'available': USE_WHISPER})

@socketio.on('reset')
def handle_reset():
    """Handle reset request"""
    global transcriptions
    transcriptions = []
    emit('update', {'text': '', 'full_text': ''}, broadcast=True)


@app.route('/upload-audio', methods=['POST'])
def upload_audio():
    """Handle audio file upload"""
    if 'audio_file' not in request.files:
        return jsonify({'error': 'No file part'}), 400

    file = request.files['audio_file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    try:
        # Save to a temporary file
        temp_audio_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(file.filename))
        file.save(temp_audio_path)

        # Transcribe audio file
        result = model.transcribe(temp_audio_path)

        # Clean up temp file
        os.unlink(temp_audio_path)

        return jsonify({'text': result['text']})

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/export', methods=['POST'])
def export_docx():
    """Export transcript to Word document"""
    data = request.get_json()
    text = data.get('text', '')
    language = data.get('language', 'auto')
    
    # Create Word document
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Add title
    title = doc.add_paragraph()
    title_run = title.add_run('Speech Transcription')
    title_run.bold = True
    title_run.font.size = Pt(14)
    
    # Add language info
    lang_names = {
        'auto': 'Auto-detect',
        'en': 'English', 
        'hi': 'Hindi',
        'kn': 'Kannada'
    }
    doc.add_paragraph(f"Language: {lang_names.get(language, language)}")
    doc.add_paragraph("")  # spacer
    
    # Add transcript text
    if text.strip():
        for para in text.split('\n'):
            if para.strip():
                doc.add_paragraph(para.strip())
    else:
        doc.add_paragraph('(No speech detected)')
    
    # Save to memory
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    
    return send_file(
        stream,
        as_attachment=True,
        download_name='transcription.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

if __name__ == '__main__':
    import os
    port = int(os.environ.get('PORT', 5000))
    print("Starting server...")
    socketio.run(app, host='0.0.0.0', port=port, debug=False, allow_unsafe_werkzeug=True)
